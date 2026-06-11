import streamlit as st
import subprocess
import os

@st.cache_resource
def install_playwright():
    subprocess.run(["playwright", "install", "chromium"])

install_playwright()

import datetime
import base64
import html
import io
import json
import re
from typing import Any, Dict, Iterable, List, Optional, Tuple

import streamlit.components.v1 as components

# ============================================================
# CONFIGURACIÓN GENERAL
# ============================================================

st.set_page_config(page_title="Generador de Procedimientos Modo 3 / LOTO", layout="wide")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&display=swap');
    html, body, [class*="css"], .stApp { font-family: 'Inter', Arial, sans-serif; }
    .main .block-container { padding-top: 1.6rem; max-width: 1320px; }
    .soft-note { font-size: 0.78rem; color: #64748B; line-height: 1.45; margin-top: -0.35rem; margin-bottom: 1rem; }
    .status-ok  { border:1px solid #FED7AA; background:#FFF7ED; color:#9A3412; border-radius:10px; padding:.8rem 1rem; font-size:.9rem; }
    .status-warn{ border:1px solid #FDE68A; background:#FFFBEB; color:#92400E; border-radius:10px; padding:.8rem 1rem; font-size:.9rem; }
</style>
""", unsafe_allow_html=True)

MODO3_COLOR     = "#FF560E"
MODO3_COLOR_HEX = "FF560E"
MODO3_FONT_HEX  = "FFFFFF"

# Colores de energías para la tabla de bloqueos
ENERGIA_COLORS: Dict[str, Tuple[str, str]] = {
    "Eléctrica":         ("000000", "FFFFFF"),
    "Neumática":         ("0284C7", "FFFFFF"),
    "Térmica":           ("DC2626", "FFFFFF"),
    "Hidráulica":        ("7C3AED", "FFFFFF"),
    "Potencial":         ("FFF200", "0F172A"),
    "Química":           ("FFF200", "0F172A"),
    "Vapor":             ("F59E0B", "111827"),
    "Agua":              ("16A34A", "FFFFFF"),
    "Amoníaco":          ("D9D9D9", "0F172A"),
    "Soda cáustica":     ("D9D9D9", "0F172A"),
    "Soda Cáustica":     ("D9D9D9", "0F172A"),
    "Ozono":             ("BAE6FD", "0F172A"),
    "Gas Carbónico":     ("C7D2FE", "111827"),
}

def _energia_colors(nombre: str) -> Tuple[str, str]:
    for key, val in ENERGIA_COLORS.items():
        if key.lower() in nombre.lower():
            return val
    return ("E5E7EB", "0F172A")

# ============================================================
# UTILIDADES DE DATOS
# ============================================================

def _html_esc(value: Any) -> str:
    if value is None:
        return ""
    return html.escape(str(value), quote=True)

def _normalize(value: Any) -> str:
    if value is None:
        return ""
    value = str(value).strip()
    return re.sub(r"\s+", " ", value)

def _upper(value: Any) -> str:
    return _normalize(value).upper()

def _get_from_dict(source: Dict[str, Any], keys: Iterable[str], default: Any = "") -> Any:
    for key in keys:
        if key in source and source.get(key) not in (None, "", [], {}):
            return source.get(key)
    return default

def _split_tasks(raw: Any) -> List[str]:
    if raw is None:
        return []
    if isinstance(raw, dict):
        for key in ("tareas", "items", "seleccionadas", "predefinidas", "values"):
            if key in raw and raw.get(key) not in (None, "", [], {}):
                return _split_tasks(raw.get(key))
        candidate = _get_from_dict(raw, ["tarea","actividad","descripcion","descripción","nombre","texto","label"], "")
        return [str(candidate).strip(" -•\t")] if str(candidate).strip() else []
    if isinstance(raw, (list, tuple)):
        tasks: List[str] = []
        for item in raw:
            if isinstance(item, dict):
                tasks.extend(_split_tasks(item))
            else:
                t = str(item).strip(" -•\t")
                if t:
                    tasks.append(t)
        return [t for t in tasks if t]
    text = str(raw).strip()
    if not text:
        return []
    return [ln.strip(" -•\t") for ln in re.split(r"[\n\r]+", text) if ln.strip(" -•\t")]

def _mode_is_modo_3(mode: Any) -> bool:
    clean = _normalize(mode).lower().replace("°", "")
    return clean in {"modo 3", "modo3", "3", "m3"}

def _auto_negocio_from_sitio(sitio: Any, default: str = "") -> str:
    s = _normalize(sitio).lower()
    if not s:
        return _normalize(default)
    if s.startswith("planta") or s.startswith("cedi"):
        return "Bebidas"
    return "Ingenios"

def _bytes_to_data_uri(content: bytes, mime: str = "image/png") -> str:
    if not content:
        return ""
    return f"data:{mime or 'image/png'};base64,{base64.b64encode(content).decode()}"

def _mime_from_filename(filename: str, default: str = "image/png") -> str:
    ext = os.path.splitext(str(filename or ""))[1].lower()
    return {".png":"image/png",".jpg":"image/jpeg",".jpeg":"image/jpeg",".webp":"image/webp"}.get(ext, default)

def _uploaded_file_data_uri(f) -> str:
    if f is None:
        return ""
    try:
        mime = getattr(f,"type","") or _mime_from_filename(getattr(f,"name",""))
        return _bytes_to_data_uri(f.getvalue(), mime)
    except Exception:
        return ""

def _local_file_data_uri(candidates: Iterable[str]) -> str:
    for c in candidates:
        if c and os.path.isfile(c):
            try:
                with open(c,"rb") as fh:
                    return _bytes_to_data_uri(fh.read(), _mime_from_filename(c))
            except Exception:
                pass
    return ""

def _logo_data_uri(uploaded=None) -> str:
    uri = _uploaded_file_data_uri(uploaded)
    if uri:
        return uri
    base = os.path.dirname(os.path.abspath(__file__)) if "__file__" in globals() else os.getcwd()
    return _local_file_data_uri([os.path.join(base,"arca.png"), os.path.join(os.getcwd(),"arca.png"),
                                  "/home/ubuntu/arca.png", "/home/ubuntu/upload/arca.png"])

def _lsr_data_uri() -> str:
    base = os.path.dirname(os.path.abspath(__file__)) if "__file__" in globals() else os.getcwd()
    return _local_file_data_uri([os.path.join(base,"LSR.png"), os.path.join(base,"lsr.png"),
                                  os.path.join(os.getcwd(),"LSR.png"), "/home/ubuntu/LSR.png"])

def _parse_date_field(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, dict):
        value = str(value.get("value") or value.get("date") or "").strip()
    value = _normalize(value)
    if not value:
        return ""
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})$", value)
    if m:
        return f"{m.group(3)}/{m.group(2)}/{m.group(1)}"
    return value

def _data_uri_to_bytes(data_uri: str) -> "tuple[bytes, str]":
    if not data_uri or not data_uri.startswith("data:"):
        return b"", ".png"
    header, _, encoded = data_uri.partition(",")
    if not encoded:
        return b"", ".png"
    mime = header.split(";",1)[0].replace("data:","").strip().lower()
    ext  = {"image/png":".png","image/jpeg":".jpg","image/jpg":".jpg","image/webp":".webp"}.get(mime,".png")
    try:
        return base64.b64decode(encoded), ext
    except Exception:
        return b"", ext

def _fecha_export_txt(value: Any) -> str:
    if isinstance(value, datetime.date):
        return value.strftime("%d/%m/%Y")
    return _normalize(value)

# ============================================================
# EXTRACCIÓN DE CONTEXTO
# ============================================================

def extract_procedure_context(payload: Dict[str, Any]) -> Dict[str, Any]:
    widgets  = payload.get("widgets")  if isinstance(payload.get("widgets"),  dict) else {}
    computed = payload.get("computed") if isinstance(payload.get("computed"), dict) else {}

    mode_final = _get_from_dict(computed, ["modo_final","modo_convalidado","modo_resultante"],
                  _get_from_dict(widgets, ["modo_final","modo_convalidado","modo_resultante"], ""))

    # Tareas: predefinidas + manuales combinadas
    manuales_raw     = _get_from_dict(widgets,  ["tareas_manuales"], "") or _get_from_dict(payload, ["tareas_manuales"], "")
    predefinidas_raw = (_get_from_dict(widgets,  ["tareas_predefinidas","tareas_predefinidas_txt","tareas_preseleccionadas","tareas_seleccionadas"], "")
                     or _get_from_dict(payload,  ["tareas_predefinidas","tareas_predefinidas_txt","tareas_preseleccionadas","tareas_seleccionadas"], ""))
    fallback_raw     = "" if (manuales_raw or predefinidas_raw) else _get_from_dict(widgets, ["tareas_txt","tareas","tareas_aplicables","actividad","actividades"], "")

    seen: set = set()
    tasks: List[str] = []
    for t in _split_tasks(predefinidas_raw) + _split_tasks(manuales_raw) + _split_tasks(fallback_raw):
        if t not in seen:
            seen.add(t); tasks.append(t)
    if not tasks:
        tasks = ["Mantenimiento, reparaciones y trabajos con exposición a energías peligrosas."]

    # Energías: dict {nombre: magnitud}
    energias_raw = _get_from_dict(computed, ["energias_seleccionadas","energias"], {})
    if not isinstance(energias_raw, dict):
        energias_raw = {}

    # Código y fecha del documento
    codigo_doc = _get_from_dict(payload, ["codigo_documento","codigo","id"],
                  _get_from_dict(widgets, ["codigo_documento","codigo","id"], ""))
    fecha_doc  = _parse_date_field(_get_from_dict(payload, ["fecha_documento","fecha"],
                  _get_from_dict(widgets, ["fecha_documento","fecha"], "")))

    sitio = _get_from_dict(widgets, ["sitio","planta","site"], "")

    return {
        "negocio":   _auto_negocio_from_sitio(sitio, _get_from_dict(widgets, ["negocio","business","unidad_negocio"], "")),
        "sitio":     sitio,
        "area":      _get_from_dict(widgets, ["area_sector","area","sector","área"], ""),
        "linea":     _get_from_dict(widgets, ["linea","línea","linea_equipo"], ""),
        "equipo":    _get_from_dict(widgets, ["equipo_desc","equipo","descripcion_equipo","nombre_equipo","maquina","máquina"], "EQUIPO SIN DESCRIPCIÓN"),
        "fabricante":_get_from_dict(widgets, ["fabricante","marca"], ""),
        "modelo":    _get_from_dict(widgets, ["modelo"], ""),
        "anio":      _get_from_dict(widgets, ["anio","año"], ""),
        "modo_inicial": _get_from_dict(computed, ["modo_inicial"], _get_from_dict(widgets, ["modo_inicial"], "")),
        "modo_final":   mode_final,
        "tareas":    tasks,
        "energias":  energias_raw,
        "codigo_documento": codigo_doc,
        "fecha_documento":  fecha_doc,
        "photo_uri": "",
    }

# ============================================================
# PUNTOS DE BLOQUEO — estructura editable
# ============================================================

def energias_to_lockpoints(energias: Dict[str, str]) -> List[Dict[str, str]]:
    """Convierte el dict de energías en una lista de puntos de bloqueo iniciales."""
    points = []
    for nombre, magnitud in energias.items():
        points.append({
            "energia":    nombre,
            "magnitud":   magnitud,
            "ubicacion":  "",
            "accion":     "",
            "verificacion": "",
            "dispositivo": "",
        })
    return points

# ============================================================
# TEXTOS ACCIÓN / VERIFICACIÓN MODO 3
# ============================================================

_ACCION_MODO3 = [
    "1. Notificar al personal afectado que el equipo será bloqueado y señalizar el área de intervención.",
    "2. Identificar todas las fuentes de energía peligrosa del equipo según la evaluación de riesgos.",
    "3. Detener el equipo mediante el control normal de PARO DE OPERACIÓN.",
    "4. Aislar cada fuente de energía accionando el dispositivo de aislamiento correspondiente (disyuntor, válvula, llave de paso, etc.).",
    "5. Aplicar el dispositivo de bloqueo (candado, grapa, cadena) en cada punto de aislamiento.",
    "6. Cada trabajador involucrado debe colocar su propio candado personal en cada punto de bloqueo.",
    "7. Cada trabajador conserva la posesión exclusiva de su llave durante toda la intervención.",
    "8. Disipar, purgar o descargar todas las energías residuales o almacenadas: presión neumática, hidráulica, vapor, resortes, partes elevadas, energía química, térmica, etc.",
    "9. Verificar la ausencia de energía intentando el arranque desde el panel de operación antes de ingresar.",
    "10. Colocar la señalización de bloqueo (tarjeta LOTO / DANGER – DO NOT OPERATE) en cada punto bloqueado.",
    "11. Confirmar que el equipo no responde a ningún comando de arranque ni puede ser energizado.",
    "12. Realizar únicamente la tarea autorizada dentro de la zona bloqueada.",
    "13. Mantener comunicación activa con todas las personas involucradas en la intervención.",
    "14. Al finalizar, retirar herramientas, piezas y residuos del interior del equipo.",
    "15. Confirmar que todo el personal se haya retirado del equipo y la zona de riesgo.",
    "16. Cada trabajador retira su propio candado y dispositivo de bloqueo.",
    "17. Restablecer los dispositivos de aislamiento en el orden correcto según el procedimiento del equipo.",
    "18. Notificar al personal afectado que el equipo será reactivado antes de energizar.",
    "19. Energizar y confirmar que el equipo opera con normalidad.",
]

_VERIFICACION_MODO3 = [
    "1. Verificar que todas las fuentes de energía hayan sido identificadas en la evaluación de riesgos del equipo.",
    "2. Verificar que exista un punto de aislamiento para cada fuente de energía identificada.",
    "3. Verificar que cada punto de aislamiento haya sido accionado y esté en posición segura.",
    "4. Verificar que cada punto de bloqueo tenga instalado al menos un candado por trabajador involucrado.",
    "5. Verificar que ningún trabajador dependa del candado de otro para su protección personal.",
    "6. Verificar que todas las energías residuales hayan sido disipadas, purgadas o descargadas.",
    "7. Verificar la ausencia de tensión eléctrica con instrumento de medición adecuado (tester, multímetro).",
    "8. Verificar la ausencia de presión neumática e hidráulica con manómetros o apertura controlada de purgas.",
    "9. Verificar la ausencia de temperatura peligrosa antes de ingresar a zonas térmicas.",
    "10. Verificar que el equipo no responda a ningún intento de arranque desde el panel.",
    "11. Verificar que la señalización LOTO esté colocada, visible e identificada en cada punto bloqueado.",
    "12. Verificar que los candados y dispositivos sean personales, únicos e intransferibles.",
    "13. Verificar que no existan bypass, puentes, anulaciones ni modificaciones a los dispositivos de bloqueo.",
    "14. Verificar que la tarea realizada sea la autorizada y no requiera acciones adicionales no previstas.",
    "15. Verificar que todas las herramientas, piezas y residuos hayan sido retirados antes del restablecimiento.",
    "16. Verificar que todo el personal haya salido del equipo antes de retirar cualquier candado.",
    "17. Verificar que cada trabajador retire únicamente su propio candado.",
    "18. Verificar que los dispositivos de aislamiento sean restablecidos en el orden correcto.",
    "19. Verificar que el equipo opere con normalidad luego del restablecimiento y que no queden condiciones inseguras.",
]

# ============================================================
# RENDER HTML
# ============================================================

def _equipment_meta(ctx: Dict[str, Any]) -> str:
    meta = []
    if ctx.get("fabricante"): meta.append(f"Fabricante: {_html_esc(ctx['fabricante'])}")
    if ctx.get("modelo"):     meta.append(f"Modelo: {_html_esc(ctx['modelo'])}")
    if ctx.get("anio"):       meta.append(f"Año: {_html_esc(ctx['anio'])}")
    return " · ".join(meta)

def _tasks_html(tasks: List[str]) -> str:
    if len(tasks) > 13:
        return f"<div class='task-inline'>{' - '.join(_html_esc(t) for t in tasks)}</div>"
    return "".join(f"<div class='task-line'>- {_html_esc(t)}</div>" for t in tasks)

def _lockpoints_table_html(points: List[Dict[str, str]]) -> str:
    """Genera la tabla Punto | Energía+Magnitud | Ubicación | Acción | Verificación | Dispositivo"""
    rows = ""
    for i, p in enumerate(points, 1):
        bg, fc = _energia_colors(p.get("energia",""))
        energia_cell = (
            f"<td style='background:#{bg};color:#{fc};font-weight:900;font-size:7.5px;"
            f"text-align:center;vertical-align:middle;padding:4px 5px;'>"
            f"{_html_esc(p.get('energia',''))}<br>"
            f"<span style='font-weight:700;font-size:7px;'>{_html_esc(p.get('magnitud',''))}</span></td>"
        )
        rows += (
            f"<tr>"
            f"<td style='text-align:center;font-weight:900;font-size:13px;background:#FFF7ED;'>{i}</td>"
            f"{energia_cell}"
            f"<td style='font-size:7.5px;vertical-align:top;padding:4px 5px;'>{_html_esc(p.get('ubicacion',''))}</td>"
            f"<td style='font-size:7.5px;vertical-align:top;padding:4px 5px;'>{_html_esc(p.get('accion',''))}</td>"
            f"<td style='font-size:7.5px;vertical-align:top;padding:4px 5px;'>{_html_esc(p.get('verificacion',''))}</td>"
            f"<td style='font-size:7.5px;vertical-align:top;padding:4px 5px;'>{_html_esc(p.get('dispositivo',''))}</td>"
            f"</tr>"
        )
    return rows

def build_modo_3_html(
    ctx: Dict[str, Any],
    lockpoints: List[Dict[str, str]],
    *,
    codigo: str,
    revision: str,
    fecha: datetime.date,
    organizacion: str,
    logo_uri: str = "",
    lsr_uri: str = "",
    personal_afectado: str,
    personal_autorizado: str,
    elaborado_por: str,
    aprobado_por: str,
    puesto_elaborado: str = "",
    puesto_aprobado: str = "",
    fecha_firma: "datetime.date | str | None" = None,
    eval_riesgos_codigo: str = "",
    eval_riesgos_fecha: str = "",
) -> str:
    fecha_txt       = fecha.strftime("%d/%m/%Y") if isinstance(fecha, datetime.date) else _normalize(fecha)
    fecha_firma_txt = fecha_firma.strftime("%d/%m/%Y") if isinstance(fecha_firma, datetime.date) else (_normalize(fecha_firma) or fecha_txt)
    task_rows       = _tasks_html(ctx.get("tareas") or [])
    photo_uri       = ctx.get("photo_uri", "")
    equipo_meta     = _equipment_meta(ctx)
    n_bloqueos      = len(lockpoints)

    logo_block  = f"<img class='logo-img' src='{logo_uri}' alt='Logo'>" if logo_uri else f"<div class='logo-text'>{_html_esc(organizacion)}</div>"
    photo_block = f"<img class='equipment-photo' src='{photo_uri}' alt='Foto'>" if photo_uri else "<div class='photo-placeholder'>FOTO DEL PASO A PASO</div>"
    lsr_block   = f"<img class='lsr-img' src='{lsr_uri}' alt='LSR'>" if lsr_uri else "<div class='lsr-placeholder'>LSR</div>"

    eval_ref = ""
    if eval_riesgos_codigo: eval_ref += f"Código: {_html_esc(eval_riesgos_codigo)}"
    if eval_riesgos_fecha:  eval_ref += f"  ·  Fecha: {_html_esc(eval_riesgos_fecha)}"
    if not eval_ref:        eval_ref = "—"

    accion_html = "".join(f"<p>{_html_esc(l)}</p>" for l in _ACCION_MODO3)
    verif_html  = "".join(f"<p>{_html_esc(l)}</p>" for l in _VERIFICACION_MODO3)
    lock_rows   = _lockpoints_table_html(lockpoints)

    C = MODO3_COLOR  # shorthand for f-strings

    return f"""<!doctype html>
<html lang="es">
<head>
<meta charset="utf-8">
<title>Procedimiento Modo 3 / LOTO - Control de Energías Peligrosas</title>
<style>
    @page {{ size: A4 portrait; margin: 5mm; }}
    * {{ box-sizing: border-box; }}
    body {{
        margin:0; background:#E5E7EB; color:#0F172A;
        font-family: Bahnschrift,'Bahnschrift SemiCondensed','Arial Narrow',Arial,sans-serif;
        font-size:8.7px; line-height:1.14;
    }}
    .sheet {{ width:760px; margin:0 auto; background:#FFFFFF; border:1.2px solid #111827; box-shadow:0 12px 28px rgba(15,23,42,.18); }}
    table {{ border-collapse:collapse; width:100%; table-layout:fixed; }}
    td, th {{ border:1px solid #111827; padding:5px 6px; vertical-align:middle; }}
    .top-black {{ background:#050505; color:#FFFFFF; font-weight:800; text-align:center; letter-spacing:.06em; font-size:11px; padding:2px 4px; text-transform:uppercase; }}
    .logo-cell {{ width:62px; text-align:center; background:#FFFFFF; padding:3px; }}
    .logo-img  {{ max-width:56px; max-height:40px; object-fit:contain; }}
    .logo-text {{ color:#B91C1C; font-size:9px; font-weight:900; line-height:1.05; text-transform:uppercase; word-break:break-word; }}
    .main-title {{ text-align:center; font-size:12.5px; font-weight:900; text-transform:uppercase; padding:5px 6px; }}
    .doc-data-label {{ background:#E5E7EB; width:68px; text-align:center; font-weight:800; text-transform:uppercase; font-size:10px; }}
    .doc-data-value {{ width:86px; text-align:center; font-weight:700; background:#F8FAFC; font-size:10px; }}
    .info-label {{ display:block; color:#334155; font-size:10px; margin-bottom:2px; }}
    .info-value {{ display:block; color:#111827; font-size:13px; font-weight:700; }}
    .info-box   {{ height:50px; text-align:center; background:#FFFFFF; }}
    .person-title {{ font-weight:900; text-align:center; background:{C}; color:#FFFFFF; font-size:8.8px; padding:3px 6px; }}
    .person-body  {{ text-align:center; min-height:28px; font-size:8.8px; white-space:pre-line; }}
    .equipment-label {{ width:62px; font-weight:800; color:#334155; background:#F8FAFC; text-align:center; }}
    .equipment-title {{ text-align:center; font-size:15px; font-weight:900; text-transform:uppercase; padding:4px 6px 2px; }}
    .equipment-meta  {{ display:block; margin-top:3px; color:#64748B; font-size:9px; font-weight:600; text-transform:none; }}
    .eval-bar {{ background:#FFFFFF; color:#0F172A; font-weight:700; font-size:8.5px; padding:4px 8px;
                 border-left:1px solid #111827; border-right:1px solid #111827; border-bottom:1px solid #111827; }}
    .orange-header th {{ background:{C}; color:#FFFFFF; font-weight:900; text-align:center; font-size:9px; }}
    .orange-bar {{ background:{C}; color:#FFFFFF; font-weight:900; text-align:center; font-size:10px;
                   border-left:1px solid #111827; border-right:1px solid #111827; padding:4px 8px; }}
    .block-zero {{ text-align:center; font-size:26px; font-weight:900; color:#000000; background:#FFF7ED; line-height:1; vertical-align:middle; }}
    .mode-box   {{ background:{C}; color:#FFFFFF; text-align:center; font-size:15px; font-weight:900; text-transform:uppercase; vertical-align:middle; }}
    .tasks-cell {{ min-height:36px; font-size:8.6px; background:#FFFFFF; padding:5px 8px; }}
    .task-line  {{ margin:0 0 4px 0; }}
    .task-inline {{ margin:0; line-height:1.35; white-space:normal; word-break:break-word; }}
    /* Tabla de bloqueos */
    .lock-table th {{ background:{C}; color:#FFFFFF; font-weight:900; font-size:8px; text-align:center; padding:3px 4px; }}
    .lock-table td {{ font-size:7.8px; vertical-align:top; padding:4px 5px; }}
    .lock-table tr:nth-child(even) td {{ background:#FFF7ED; }}
    .photo-area {{ height:210px; background:#FFFFFF; text-align:center; vertical-align:middle; padding:7px; }}
    .equipment-photo {{ max-width:100%; max-height:195px; object-fit:contain; border:1px solid #CBD5E1; background:#F8FAFC; padding:3px; }}
    .photo-placeholder {{ height:190px; display:flex; align-items:center; justify-content:center;
                          color:#94A3B8; font-size:14px; font-weight:800; letter-spacing:.03em;
                          border:1px dashed #CBD5E1; background:#F8FAFC; }}
    .dark-head th {{ background:#3B3B3B; color:#FFFFFF; font-weight:900; text-align:center; font-size:8.3px; padding:3px 2px; }}
    .procedure-table td {{ font-size:7.8px; text-align:left; background:#FFFFFF; vertical-align:top; padding:7px 9px; }}
    .procedure-note p {{ margin:0 0 4px 0; font-size:8.4px; line-height:1.42; }}
    .procedure-note p:last-child {{ margin-bottom:0; }}
    .legend-title {{ background:{C}; color:#FFFFFF; font-weight:900; text-align:center; padding:2px; }}
    .energy-legend td,.lock-legend td {{ text-align:center; font-weight:900; font-size:7.6px; color:#0F172A; padding:3px 2px; }}
    .e-electric  {{ background:#000000; color:#FFFFFF !important; }}
    .e-neumatic  {{ background:#0284C7; color:#FFFFFF !important; }}
    .e-amoniaco  {{ background:linear-gradient(90deg,#D9D9D9 0% 15%,#F59E0B 15% 26%,#D9D9D9 26% 42%,#F59E0B 42% 53%,#D9D9D9 53% 68%,#F59E0B 68% 79%,#D9D9D9 79% 100%); }}
    .e-termica   {{ background:#DC2626; color:#FFFFFF !important; }}
    .e-hidraulica{{ background:#7C3AED; color:#FFFFFF !important; }}
    .e-potencial {{ background:linear-gradient(90deg,#FFF200 0% 22%,#050505 22% 34%,#FFF200 34% 62%,#050505 62% 74%,#FFF200 74% 100%); }}
    .e-quimica   {{ background:#FFF200; }}
    .e-vapor     {{ background:#F59E0B; color:#111827 !important; }}
    .e-agua      {{ background:#16A34A; color:#FFFFFF !important; }}
    .e-soda      {{ background:linear-gradient(90deg,#D9D9D9 0% 25%,#F59E0B 25% 37%,#D9D9D9 37% 60%,#F59E0B 60% 72%,#D9D9D9 72% 100%); }}
    .e-ozono     {{ background:#BAE6FD; }}
    .e-gas       {{ background:#C7D2FE; color:#111827 !important; }}
    .l-mmto      {{ background:#EF0000; color:#FFFFFF !important; }}
    .l-calidad   {{ background:#FFF200; }}
    .l-produccion{{ background:#16A34A; color:#FFFFFF !important; }}
    .l-edilicio  {{ background:#0F7DBD; color:#FFFFFF !important; }}
    .l-supervisor{{ background:#050505; color:#FFFFFF !important; }}
    .lock-legend {{ table-layout:fixed; }}
    .lock-legend td {{ width:20%; height:24px; line-height:1.2; vertical-align:middle; overflow:hidden; font-weight:900; font-size:7.6px; }}
    .lsr-section {{ border-left:1px solid #111827; border-right:1px solid #111827; border-bottom:1px solid #111827; text-align:center; padding:8px 6px; background:#FFFFFF; }}
    .lsr-img  {{ max-width:100%; max-height:160px; object-fit:contain; }}
    .lsr-placeholder {{ height:80px; display:flex; align-items:center; justify-content:center; color:#94A3B8; font-size:13px; font-weight:800; border:1px dashed #CBD5E1; background:#F8FAFC; }}
    .signature-cell {{ height:30px; font-size:7.8px; color:#334155; background:#F8FAFC; text-align:center; }}
    .small-muted {{ color:#64748B; font-size:8.8px; font-weight:600; }}
    @media print {{ body{{ background:#FFFFFF; }} .sheet{{ box-shadow:none; margin:0; width:100%; }} }}
</style>
</head>
<body>
<div class="sheet">

<!-- ENCABEZADO -->
<table>
  <tr>
    <td class="logo-cell" rowspan="3">{logo_block}</td>
    <td class="top-black" colspan="2">CONTROL DE ENERGÍAS PELIGROSAS</td>
    <td class="doc-data-label">Código</td>
    <td class="doc-data-value">{_html_esc(codigo)}</td>
  </tr>
  <tr>
    <td class="main-title" colspan="2" rowspan="2">PROCEDIMIENTO ESPECÍFICO PARA CONTROL DE ENERGÍAS</td>
    <td class="doc-data-label">Revisión</td>
    <td class="doc-data-value">{_html_esc(revision)}</td>
  </tr>
  <tr>
    <td class="doc-data-label">Fecha</td>
    <td class="doc-data-value">{_html_esc(fecha_txt)}</td>
  </tr>
</table>

<!-- INFO SITIO + PERSONAL -->
<table>
  <tr>
    <td class="info-box" style="width:62px;"><span class="info-label">Negocio:</span><span class="info-value">{_html_esc(ctx.get('negocio') or '-')}</span></td>
    <td class="info-box" style="width:86px;"><span class="info-label">Sitio:</span><span class="info-value">{_html_esc(ctx.get('sitio') or '-')}</span></td>
    <td class="info-box" style="width:86px;"><span class="info-label">Área:</span><span class="info-value">{_html_esc(ctx.get('area') or '-')}</span></td>
    <td class="info-box" style="width:112px;"><span class="info-label">Línea:</span><span class="info-value">{_html_esc(ctx.get('linea') or '-')}</span></td>
    <td style="padding:0;">
      <table>
        <tr><td class="person-title">Personal afectado - Puestos de trabajo</td></tr>
        <tr><td class="person-body">{_html_esc(personal_afectado)}</td></tr>
        <tr><td class="person-title">Personal autorizado - Puestos de trabajo</td></tr>
        <tr><td class="person-body">{_html_esc(personal_autorizado)}</td></tr>
      </table>
    </td>
  </tr>
</table>

<!-- EQUIPO -->
<table>
  <tr>
    <td class="equipment-label">Equipo:</td>
    <td class="equipment-title">{_html_esc(ctx.get('equipo') or 'EQUIPO')}<span class="equipment-meta">{equipo_meta}</span></td>
  </tr>
</table>

<!-- EVALUACIÓN DE RIESGOS -->
<div class="eval-bar"><strong>Evaluación de riesgos de referencia:</strong>&nbsp;&nbsp;{eval_ref}</div>

<!-- TAREAS -->
<table>
  <tr class="orange-header">
    <th style="width:80px;">Puntos de<br>Bloqueo</th>
    <th style="width:90px;">Modo de<br>Intervención</th>
    <th>Listado de tareas aplicable al presente procedimiento</th>
  </tr>
  <tr>
    <td class="block-zero">{n_bloqueos}</td>
    <td class="mode-box">MODO 3<br>LOTO</td>
    <td class="tasks-cell">{task_rows}</td>
  </tr>
</table>

<!-- BARRA -->
<div class="orange-bar">Procedimiento - Control de Energías Peligrosas / LOTO</div>

<!-- FOTO -->
<table><tr><td class="photo-area">{photo_block}</td></tr></table>

<!-- TABLA DE PUNTOS DE BLOQUEO -->
<table class="lock-table">
  <tr>
    <th style="width:36px;">Punto<br>N°</th>
    <th style="width:110px;">Fuente de Energía<br>y Magnitud</th>
    <th style="width:130px;">Ubicación</th>
    <th>Acción</th>
    <th>Verificación</th>
    <th style="width:110px;">Dispositivo de<br>Bloqueo</th>
  </tr>
  {lock_rows}
</table>

<!-- PROCEDIMIENTO: Acción y Verificación -->
<table class="procedure-table">
  <tr class="dark-head">
    <th style="width:50%;">Acción</th>
    <th style="width:50%;">Verificación</th>
  </tr>
  <tr>
    <td class="procedure-note">{accion_html}</td>
    <td class="procedure-note">{verif_html}</td>
  </tr>
</table>

<!-- LEYENDA ENERGÍAS -->
<div class="legend-title">Clasificación de Energías Peligrosas</div>
<table class="energy-legend">
  <tr>
    <td class="e-electric">E: Eléctrica</td><td class="e-neumatic">N: Neumática</td>
    <td class="e-amoniaco">AM: Amoníaco</td><td class="e-termica">T: Térmica</td>
    <td class="e-hidraulica">H: Hidráulica</td><td class="e-potencial">P: Potencial</td>
  </tr>
  <tr>
    <td class="e-quimica">Q: Química</td><td class="e-vapor">V: Vapor</td>
    <td class="e-agua">A: Agua</td><td class="e-soda">SC: Soda Cáustica</td>
    <td class="e-ozono">Oz: Ozono</td><td class="e-gas">GC: Gas Carbónico</td>
  </tr>
</table>

<!-- LEYENDA CANDADOS -->
<div class="legend-title">Clasificación de Candados según sector y función</div>
<table class="lock-legend">
  <tr>
    <td class="l-mmto">Mantenimiento<br>Industrial</td>
    <td class="l-calidad">Calidad</td>
    <td class="l-produccion">Producción</td>
    <td class="l-edilicio">Mantenimiento Edilicio<br>Contratistas</td>
    <td class="l-supervisor">Supervisor MMTO Industrial<br>(Bloqueo Departamental)</td>
  </tr>
</table>

<!-- LSR -->
<div class="lsr-section">{lsr_block}</div>

<!-- FIRMAS -->
<table>
  <tr>
    <td class="signature-cell">
      Elaborado por: <strong>{_html_esc(elaborado_por or '-')}</strong><br>
      <span class="small-muted">Puesto: {_html_esc(puesto_elaborado or '-')} · Fecha: {_html_esc(fecha_firma_txt)}</span>
    </td>
    <td class="signature-cell">
      Aprobado por: <strong>{_html_esc(aprobado_por or '-')}</strong><br>
      <span class="small-muted">Puesto: {_html_esc(puesto_aprobado or '-')} · Fecha: {_html_esc(fecha_firma_txt)}</span>
    </td>
  </tr>
</table>

</div></body></html>"""

# ============================================================
# EXPORTACIÓN WORD
# ============================================================

def _docx_set_cell_shading(cell, fill):
    from docx.oxml import OxmlElement; from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#","").upper())

def _docx_set_cell_borders(cell, color="111827", size="6"):
    from docx.oxml import OxmlElement; from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_b  = tc_pr.first_child_found_in("w:tcBorders")
    if tc_b is None:
        tc_b = OxmlElement("w:tcBorders"); tc_pr.append(tc_b)
    for edge in ("top","left","bottom","right"):
        el = tc_b.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}"); tc_b.append(el)
        el.set(qn("w:val"),"single"); el.set(qn("w:sz"),size)
        el.set(qn("w:space"),"0");   el.set(qn("w:color"),color.replace("#","").upper())

def _docx_set_cell_margins(cell, top=45, start=45, bottom=45, end=45):
    from docx.oxml import OxmlElement; from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_m  = tc_pr.first_child_found_in("w:tcMar")
    if tc_m is None:
        tc_m = OxmlElement("w:tcMar"); tc_pr.append(tc_m)
    for m, v in {"top":top,"start":start,"bottom":bottom,"end":end}.items():
        n = tc_m.find(qn(f"w:{m}"))
        if n is None:
            n = OxmlElement(f"w:{m}"); tc_m.append(n)
        n.set(qn("w:w"), str(v)); n.set(qn("w:type"),"dxa")

def _docx_set_cell_width(cell, width_cm):
    from docx.oxml import OxmlElement; from docx.oxml.ns import qn; from docx.shared import Cm
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w  = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm*567))); tc_w.set(qn("w:type"),"dxa")

def _docx_clear_cell(cell):
    cell.text = ""

def _docx_write_cell(cell, text="", *, bold=False, size=7.5, color="0F172A", fill=None, align="center", valign="center"):
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT; from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    _docx_clear_cell(cell)
    if fill: _docx_set_cell_shading(cell, fill)
    _docx_set_cell_borders(cell); _docx_set_cell_margins(cell)
    cell.vertical_alignment = {"top":WD_CELL_VERTICAL_ALIGNMENT.TOP,"bottom":WD_CELL_VERTICAL_ALIGNMENT.BOTTOM}.get(valign, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
    para = cell.paragraphs[0]
    para.alignment = {"left":WD_ALIGN_PARAGRAPH.LEFT,"right":WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    para.paragraph_format.space_before = 0; para.paragraph_format.space_after = 0
    run = para.add_run(str(text or ""))
    run.bold = bold; run.font.name = "Bahnschrift"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color.replace("#","").upper())

def _docx_apply_table_grid(table, widths_cm=None):
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE; from docx.shared import Cm
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            _docx_set_cell_borders(cell); _docx_set_cell_margins(cell)
            if widths_cm and idx < len(widths_cm): _docx_set_cell_width(cell, widths_cm[idx])
    if widths_cm:
        for idx, w in enumerate(widths_cm):
            try: table.columns[idx].width = Cm(w)
            except: pass

def html_to_word_bytes(
    ctx, lockpoints, *, codigo, revision, fecha, organizacion,
    logo_uri="", lsr_uri="", personal_afectado, personal_autorizado,
    elaborado_por, aprobado_por, puesto_elaborado="", puesto_aprobado="",
    fecha_firma=None, eval_riesgos_codigo="", eval_riesgos_fecha="",
) -> bytes:
    from docx import Document; from docx.enum.section import WD_SECTION_START
    from docx.enum.text import WD_ALIGN_PARAGRAPH; from docx.shared import Cm, Pt, RGBColor
    from docx.oxml.ns import qn as _qn

    fecha_txt       = _fecha_export_txt(fecha)
    fecha_firma_txt = _fecha_export_txt(fecha_firma) or fecha_txt
    equipo_meta     = _equipment_meta(ctx)
    n_bloqueos      = len(lockpoints)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(0.45); sec.bottom_margin = Cm(0.45)
    sec.left_margin = Cm(0.45); sec.right_margin = Cm(0.45)
    doc.styles["Normal"].font.name = "Bahnschrift"
    doc.styles["Normal"].font.size = Pt(7.5)

    TW = 17.8  # total usable width cm

    # Encabezado
    hdr = doc.add_table(rows=3, cols=5)
    _docx_apply_table_grid(hdr, [2.0, 5.0, 5.2, 2.4, 3.2])
    hdr.cell(0,0).merge(hdr.cell(2,0)); hdr.cell(0,1).merge(hdr.cell(0,2)); hdr.cell(1,1).merge(hdr.cell(2,2))
    logo_cell = hdr.cell(0,0); _docx_write_cell(logo_cell,"",fill="FFFFFF")
    logo_bytes, _ = _data_uri_to_bytes(logo_uri)
    if logo_bytes:
        try:
            p = logo_cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(logo_bytes), width=Cm(1.55))
        except: _docx_write_cell(logo_cell, organizacion, bold=True, size=7.0, color="B91C1C", fill="FFFFFF")
    else: _docx_write_cell(logo_cell, organizacion, bold=True, size=7.0, color="B91C1C", fill="FFFFFF")
    _docx_write_cell(hdr.cell(0,1),"CONTROL DE ENERGÍAS PELIGROSAS",bold=True,size=8.5,color="FFFFFF",fill="050505")
    _docx_write_cell(hdr.cell(1,1),"PROCEDIMIENTO ESPECÍFICO PARA CONTROL DE ENERGÍAS",bold=True,size=9.0,fill="FFFFFF")
    for ri, lbl, val in ((0,"Código",codigo),(1,"Revisión",revision),(2,"Fecha",fecha_txt)):
        _docx_write_cell(hdr.cell(ri,3),lbl,bold=True,size=7.5,fill="E5E7EB")
        _docx_write_cell(hdr.cell(ri,4),val,bold=True,size=7.5,fill="F8FAFC")

    # Info sitio + personal
    info = doc.add_table(rows=1, cols=5)
    _docx_apply_table_grid(info, [2.0,2.8,2.8,3.6,6.6])
    for idx,(lbl,val) in enumerate([("Negocio:",ctx.get("negocio") or "-"),("Sitio:",ctx.get("sitio") or "-"),("Área:",ctx.get("area") or "-"),("Línea:",ctx.get("linea") or "-")]):
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        cell = info.cell(0,idx); _docx_clear_cell(cell); _docx_set_cell_shading(cell,"FFFFFF")
        _docx_set_cell_borders(cell); _docx_set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = 0; para.paragraph_format.space_after = 0
        r1 = para.add_run(lbl+"\n"); r1.bold=False; r1.font.name="Bahnschrift"; r1.font.size=Pt(7.2); r1.font.color.rgb=RGBColor.from_string("334155")
        r2 = para.add_run(str(val)); r2.bold=True;  r2.font.name="Bahnschrift"; r2.font.size=Pt(9.0);  r2.font.color.rgb=RGBColor.from_string("111827")
    person_cell = info.cell(0,4); _docx_clear_cell(person_cell)
    _docx_set_cell_borders(person_cell); _docx_set_cell_margins(person_cell,top=0,start=0,bottom=0,end=0)
    for p_elem in person_cell._tc.findall(_qn("w:p")): person_cell._tc.remove(p_elem)
    nested = person_cell.add_table(rows=4,cols=1); _docx_apply_table_grid(nested,[6.4])
    _docx_write_cell(nested.cell(0,0),"Personal afectado - Puestos de trabajo",bold=True,size=6.8,color=MODO3_FONT_HEX,fill=MODO3_COLOR_HEX)
    _docx_write_cell(nested.cell(1,0),personal_afectado,size=6.8,fill="FFFFFF")
    _docx_write_cell(nested.cell(2,0),"Personal autorizado - Puestos de trabajo",bold=True,size=6.8,color=MODO3_FONT_HEX,fill=MODO3_COLOR_HEX)
    _docx_write_cell(nested.cell(3,0),personal_autorizado,size=6.8,fill="FFFFFF")

    # Equipo
    eq = doc.add_table(rows=1,cols=2); _docx_apply_table_grid(eq,[2.0,15.8])
    _docx_write_cell(eq.cell(0,0),"Equipo:",bold=True,size=7.4,fill="F8FAFC")
    _docx_write_cell(eq.cell(0,1),f"{_upper(ctx.get('equipo') or 'EQUIPO')}\n{equipo_meta}",bold=True,size=9.0,fill="FFFFFF")

    # Evaluación riesgos
    ev = doc.add_table(rows=1,cols=1); _docx_apply_table_grid(ev,[TW])
    ev_txt = "Evaluación de riesgos de referencia:  "
    if eval_riesgos_codigo: ev_txt += f"Código: {eval_riesgos_codigo}"
    if eval_riesgos_fecha:  ev_txt += f"  ·  Fecha: {eval_riesgos_fecha}"
    if ev_txt.strip().endswith(":"): ev_txt += " —"
    _docx_write_cell(ev.cell(0,0),ev_txt,bold=True,size=7.0,color="0F172A",fill="FFFFFF",align="left")

    # Tareas
    tt = doc.add_table(rows=2,cols=3); _docx_apply_table_grid(tt,[2.0,2.8,13.0])
    for i,t in enumerate(["Puntos de\nBloqueo","Modo de\nIntervención","Listado de tareas aplicable al presente procedimiento"]):
        _docx_write_cell(tt.cell(0,i),t,bold=True,size=7.2,color=MODO3_FONT_HEX,fill=MODO3_COLOR_HEX)
    _docx_write_cell(tt.cell(1,0),str(n_bloqueos),bold=True,size=22.0,fill="FFF7ED")
    _docx_write_cell(tt.cell(1,1),"MODO 3\nLOTO",bold=True,size=11.0,color=MODO3_FONT_HEX,fill=MODO3_COLOR_HEX)
    _tareas = ctx.get("tareas") or []
    tasks_str = " - ".join(_tareas) if len(_tareas) > 13 else "\n".join(f"- {t}" for t in _tareas)
    _docx_write_cell(tt.cell(1,2),tasks_str,size=6.8,fill="FFFFFF",align="left")

    # Barra
    bar = doc.add_table(rows=1,cols=1); _docx_apply_table_grid(bar,[TW])
    _docx_write_cell(bar.cell(0,0),"Procedimiento - Control de Energías Peligrosas / LOTO",bold=True,size=7.6,color=MODO3_FONT_HEX,fill=MODO3_COLOR_HEX)

    # Foto
    ph = doc.add_table(rows=1,cols=1); _docx_apply_table_grid(ph,[TW])
    ph_cell = ph.cell(0,0); _docx_write_cell(ph_cell,"FOTO DEL PASO A PASO",bold=True,size=11.0,color="94A3B8",fill="FFFFFF")
    photo_bytes, _ = _data_uri_to_bytes(ctx.get("photo_uri",""))
    if photo_bytes:
        try:
            _docx_clear_cell(ph_cell); p = ph_cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(photo_bytes), width=Cm(12.0)); _docx_set_cell_borders(ph_cell)
        except: _docx_write_cell(ph_cell,"FOTO DEL PASO A PASO",bold=True,size=11.0,color="94A3B8",fill="FFFFFF")

    # Tabla de puntos de bloqueo
    lk_hdr = doc.add_table(rows=1,cols=1); _docx_apply_table_grid(lk_hdr,[TW])
    _docx_write_cell(lk_hdr.cell(0,0),"Puntos de Bloqueo LOTO",bold=True,size=7.5,color=MODO3_FONT_HEX,fill=MODO3_COLOR_HEX)

    COL_W = [1.0, 3.0, 3.0, 3.6, 3.6, 3.6]
    lk = doc.add_table(rows=1+len(lockpoints), cols=6)
    _docx_apply_table_grid(lk, COL_W)
    for i,h in enumerate(["N°","Fuente de Energía\ny Magnitud","Ubicación","Acción","Verificación","Dispositivo de Bloqueo"]):
        _docx_write_cell(lk.cell(0,i),h,bold=True,size=6.5,color=MODO3_FONT_HEX,fill=MODO3_COLOR_HEX)
    for ri, p in enumerate(lockpoints, 1):
        bg, fc = _energia_colors(p.get("energia",""))
        _docx_write_cell(lk.cell(ri,0),str(ri),bold=True,size=10.0,fill="FFF7ED")
        _docx_write_cell(lk.cell(ri,1),f"{p.get('energia','')}\n{p.get('magnitud','')}",bold=True,size=6.5,color=fc,fill=bg)
        _docx_write_cell(lk.cell(ri,2),p.get("ubicacion",""),size=6.5,fill="FFFFFF",align="left",valign="top")
        _docx_write_cell(lk.cell(ri,3),p.get("accion",""),size=6.5,fill="FFFFFF",align="left",valign="top")
        _docx_write_cell(lk.cell(ri,4),p.get("verificacion",""),size=6.5,fill="FFFFFF",align="left",valign="top")
        _docx_write_cell(lk.cell(ri,5),p.get("dispositivo",""),size=6.5,fill="FFFFFF",align="left",valign="top")

    # Procedimiento Acción / Verificación
    pr = doc.add_table(rows=2,cols=2); _docx_apply_table_grid(pr,[8.9,8.9])
    _docx_write_cell(pr.cell(0,0),"Acción",bold=True,size=7.0,color="FFFFFF",fill="3B3B3B")
    _docx_write_cell(pr.cell(0,1),"Verificación",bold=True,size=7.0,color="FFFFFF",fill="3B3B3B")
    _docx_write_cell(pr.cell(1,0),"\n\n".join(_ACCION_MODO3),size=6.5,fill="FFFFFF",align="left",valign="top")
    _docx_write_cell(pr.cell(1,1),"\n\n".join(_VERIFICACION_MODO3),size=6.5,fill="FFFFFF",align="left",valign="top")

    # Leyenda energías
    le = doc.add_table(rows=1,cols=1); _docx_apply_table_grid(le,[TW])
    _docx_write_cell(le.cell(0,0),"Clasificación de Energías Peligrosas",bold=True,size=7.0,color=MODO3_FONT_HEX,fill=MODO3_COLOR_HEX)
    en = doc.add_table(rows=2,cols=6); _docx_apply_table_grid(en,[TW/6]*6)
    energy_items = [
        (0,0,"E: Eléctrica","000000","FFFFFF",None),(0,1,"N: Neumática","0284C7","FFFFFF",None),
        (0,2,"AM: Amoníaco","D9D9D9","0F172A",("D9D9D9","F59E0B",[(15,26),(42,53),(68,79)])),
        (0,3,"T: Térmica","DC2626","FFFFFF",None),(0,4,"H: Hidráulica","7C3AED","FFFFFF",None),
        (0,5,"P: Potencial","FFF200","555555",("FFF200","050505",[(22,34),(62,74)])),
        (1,0,"Q: Química","FFF200","0F172A",None),(1,1,"V: Vapor","F59E0B","111827",None),
        (1,2,"A: Agua","16A34A","FFFFFF",None),
        (1,3,"SC: Soda Cáustica","D9D9D9","0F172A",("D9D9D9","F59E0B",[(25,37),(60,72)])),
        (1,4,"Oz: Ozono","BAE6FD","0F172A",None),(1,5,"GC: Gas Carbónico","C7D2FE","111827",None),
    ]
    for row,col,lbl,fill,fc,stripe in energy_items:
        cell = en.cell(row,col)
        if stripe:
            try:
                from PIL import Image as _PIL, ImageDraw as _Draw
                bg_s, sc_s, pos = stripe
                img = _PIL.new("RGB",(120,30),tuple(int(bg_s[i:i+2],16) for i in (0,2,4)))
                draw = _Draw.Draw(img)
                sc_rgb = tuple(int(sc_s[i:i+2],16) for i in (0,2,4))
                for x0p,x1p in pos:
                    draw.rectangle([int(120*x0p/100),0,int(120*x1p/100)-1,29],fill=sc_rgb)
                buf = io.BytesIO(); img.save(buf,format="PNG"); buf.seek(0)
                # simplified: just use solid color for stripe cells in Word
                _docx_write_cell(cell,lbl,bold=True,size=5.8,color=fc,fill=bg_s)
            except: _docx_write_cell(cell,lbl,bold=True,size=5.8,color=fc,fill=fill)
        else: _docx_write_cell(cell,lbl,bold=True,size=5.8,color=fc,fill=fill)

    # Leyenda candados
    lc = doc.add_table(rows=1,cols=1); _docx_apply_table_grid(lc,[TW])
    _docx_write_cell(lc.cell(0,0),"Clasificación de Candados según sector y función",bold=True,size=7.0,color=MODO3_FONT_HEX,fill=MODO3_COLOR_HEX)
    lck = doc.add_table(rows=1,cols=5); _docx_apply_table_grid(lck,[TW/5]*5)
    for i,(lbl,fill,fc) in enumerate([("Mantenimiento\nIndustrial","EF0000","FFFFFF"),("Calidad","FFF200","0F172A"),("Producción","16A34A","FFFFFF"),("Mantenimiento Edilicio\nContratistas","0F7DBD","FFFFFF"),("Supervisor MMTO Industrial\n(Bloqueo Departamental)","050505","FFFFFF")]):
        _docx_write_cell(lck.cell(0,i),lbl,bold=True,size=5.8,color=fc,fill=fill)

    # LSR
    lsr_t = doc.add_table(rows=1,cols=1); _docx_apply_table_grid(lsr_t,[TW])
    lsr_cell = lsr_t.cell(0,0); _docx_write_cell(lsr_cell,"LSR",bold=True,size=10.0,color="94A3B8",fill="FFFFFF")
    lsr_b, _ = _data_uri_to_bytes(lsr_uri)
    if lsr_b:
        try:
            _docx_clear_cell(lsr_cell); p = lsr_cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(lsr_b), width=Cm(15.0)); _docx_set_cell_borders(lsr_cell)
        except: _docx_write_cell(lsr_cell,"LSR",bold=True,size=10.0,color="94A3B8",fill="FFFFFF")

    # Firmas
    ft = doc.add_table(rows=1,cols=2); _docx_apply_table_grid(ft,[8.9,8.9])
    _docx_write_cell(ft.cell(0,0),f"Elaborado por: {elaborado_por or '-'}\nPuesto: {puesto_elaborado or '-'} · Fecha: {fecha_firma_txt}",size=6.4,fill="F8FAFC")
    _docx_write_cell(ft.cell(0,1),f"Aprobado por: {aprobado_por or '-'}\nPuesto: {puesto_aprobado or '-'} · Fecha: {fecha_firma_txt}",size=6.4,fill="F8FAFC")

    for para in doc.paragraphs:
        para.paragraph_format.space_before = 0; para.paragraph_format.space_after = 0

    out = io.BytesIO(); doc.save(out); return out.getvalue()

# ============================================================
# EXPORTACIÓN EXCEL
# ============================================================

def _xlsx_style(ws, rng, fill=None, font_color="0F172A", bold=False, size=8.0, align="center"):
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    side = Side(style="thin",color="111827")
    border = Border(left=side,right=side,top=side,bottom=side)
    for row in ws[rng]:
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(horizontal=align,vertical="center",wrap_text=True)
            cell.font = Font(name="Bahnschrift",size=size,bold=bold,color=font_color)
            if fill: cell.fill = PatternFill("solid",fgColor=fill.replace("#",""))

def _xlsx_mw(ws, rng, value, *, fill=None, font_color="0F172A", bold=False, size=8.0, align="center"):
    ws.merge_cells(rng); ws[rng.split(":")[0]] = value
    _xlsx_style(ws, rng, fill=fill, font_color=font_color, bold=bold, size=size, align=align)

def _xlsx_add_image(ws, data_uri, anchor, *, max_width_px, max_height_px):
    if not data_uri: return
    try:
        from openpyxl.drawing.image import Image as XLImage; from PIL import Image as PILImage
        ib, _ = _data_uri_to_bytes(data_uri)
        if not ib: return
        pil = PILImage.open(io.BytesIO(ib)).convert("RGBA"); pil.thumbnail((max_width_px,max_height_px),PILImage.LANCZOS)
        s = io.BytesIO(); pil.save(s,format="PNG"); s.seek(0)
        xl = XLImage(s); xl.anchor = anchor; ws.add_image(xl)
    except: pass

def build_modo_3_excel_bytes(
    ctx, lockpoints, *, codigo, revision, fecha, organizacion,
    logo_uri="", lsr_uri="", personal_afectado, personal_autorizado,
    elaborado_por, aprobado_por, puesto_elaborado="", puesto_aprobado="",
    fecha_firma=None, eval_riesgos_codigo="", eval_riesgos_fecha="",
) -> bytes:
    from openpyxl import Workbook; from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    fecha_txt       = _fecha_export_txt(fecha)
    fecha_firma_txt = _fecha_export_txt(fecha_firma) or fecha_txt
    equipo_meta     = _equipment_meta(ctx)
    n_bloqueos      = len(lockpoints)

    wb = Workbook(); ws = wb.active; ws.title = "Modo 3 LOTO"
    ws.sheet_view.showGridLines = False
    ws.page_setup.orientation = "portrait"; ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToWidth = 1; ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    for m in ("left","right","top","bottom","header","footer"):
        setattr(ws.page_margins, m, 0.18 if m not in ("header","footer") else 0)

    for col in range(1,25): ws.column_dimensions[get_column_letter(col)].width = 4.2
    for row in range(1,80): ws.row_dimensions[row].height = 18

    white = PatternFill("solid",fgColor="FFFFFF")
    for r in range(1,80):
        for c in range(1,25):
            ws.cell(row=r,column=c).fill = white
            ws.cell(row=r,column=c).font = Font(name="Bahnschrift",size=8)

    # Encabezado
    _xlsx_mw(ws,"A1:C3","" if logo_uri else organizacion,fill="FFFFFF",font_color="B91C1C",bold=True,size=8)
    _xlsx_add_image(ws,logo_uri,"A1",max_width_px=78,max_height_px=52)
    _xlsx_mw(ws,"D1:P1","CONTROL DE ENERGÍAS PELIGROSAS",fill="050505",font_color="FFFFFF",bold=True,size=9)
    _xlsx_mw(ws,"D2:P3","PROCEDIMIENTO ESPECÍFICO PARA CONTROL DE ENERGÍAS",fill="FFFFFF",bold=True,size=10)
    for rng,lbl,val in (("Q1:S1","Código",codigo),("Q2:S2","Revisión",revision),("Q3:S3","Fecha",fecha_txt)):
        _xlsx_mw(ws,rng,lbl,fill="E5E7EB",bold=True,size=8)
        _xlsx_mw(ws,rng.replace("Q","T").replace("S","X"),val,fill="F8FAFC",bold=True,size=8)

    # Sitio + personal
    def _lv(ws,rng,lbl,val):
        from openpyxl.styles import Alignment,Border,Font,PatternFill,Side
        from openpyxl.cell.rich_text import CellRichText,TextBlock; from openpyxl.styles.fonts import Font as RF
        tl = rng.split(":")[0]; ws.merge_cells(rng)
        side = Side(style="thin",color="111827"); border = Border(left=side,right=side,top=side,bottom=side)
        try:
            ws[tl] = CellRichText(TextBlock(RF(name="Bahnschrift",size=8,bold=False,color="FF334155"),lbl+"\n"),
                                   TextBlock(RF(name="Bahnschrift",size=10,bold=True,color="FF111827"),val))
        except: ws[tl] = f"{lbl}\n{val}"
        ws[tl].alignment = Alignment(horizontal="center",vertical="center",wrap_text=True)
        ws[tl].border = border; ws[tl].fill = PatternFill("solid",fgColor="FFFFFF")

    _lv(ws,"A4:C6","Negocio:",ctx.get("negocio") or "-"); _lv(ws,"D4:F6","Sitio:",ctx.get("sitio") or "-")
    _lv(ws,"G4:I6","Área:",ctx.get("area") or "-");       _lv(ws,"J4:M6","Línea:",ctx.get("linea") or "-")
    _xlsx_mw(ws,"N4:X4","Personal afectado - Puestos de trabajo",fill=MODO3_COLOR_HEX,font_color=MODO3_FONT_HEX,bold=True,size=7)
    _xlsx_mw(ws,"N5:X5",personal_afectado,fill="FFFFFF",size=7)
    _xlsx_mw(ws,"N6:X6","Personal autorizado - Puestos de trabajo",fill=MODO3_COLOR_HEX,font_color=MODO3_FONT_HEX,bold=True,size=7)
    _xlsx_mw(ws,"N7:X7",personal_autorizado,fill="FFFFFF",size=7)

    # Equipo
    _xlsx_mw(ws,"A8:C9","Equipo:",fill="F8FAFC",bold=True,size=8)
    _xlsx_mw(ws,"D8:X9",f"{_upper(ctx.get('equipo') or 'EQUIPO')}\n{equipo_meta}",fill="FFFFFF",bold=True,size=11)

    # Evaluación riesgos
    ev_txt = "Evaluación de riesgos de referencia:  "
    if eval_riesgos_codigo: ev_txt += f"Código: {eval_riesgos_codigo}"
    if eval_riesgos_fecha:  ev_txt += f"  ·  Fecha: {eval_riesgos_fecha}"
    _xlsx_mw(ws,"A10:X10",ev_txt,fill="FFFFFF",font_color="0F172A",bold=True,size=7,align="left")

    # Tareas
    _xlsx_mw(ws,"A11:C11","Puntos de\nBloqueo",fill=MODO3_COLOR_HEX,font_color=MODO3_FONT_HEX,bold=True,size=7)
    _xlsx_mw(ws,"D11:F11","Modo de\nIntervención",fill=MODO3_COLOR_HEX,font_color=MODO3_FONT_HEX,bold=True,size=7)
    _xlsx_mw(ws,"G11:X11","Listado de tareas aplicable al presente procedimiento",fill=MODO3_COLOR_HEX,font_color=MODO3_FONT_HEX,bold=True,size=7)
    _xlsx_mw(ws,"A12:C15",str(n_bloqueos),fill="FFF7ED",bold=True,size=22)
    _xlsx_mw(ws,"D12:F15","MODO 3\nLOTO",fill=MODO3_COLOR_HEX,font_color=MODO3_FONT_HEX,bold=True,size=11)
    _tareas = ctx.get("tareas") or []
    tasks_str = " - ".join(_tareas) if len(_tareas) > 13 else "\n".join(f"- {t}" for t in _tareas)
    _xlsx_mw(ws,"G12:X15",tasks_str,fill="FFFFFF",size=7,align="left")

    # Barra + foto
    _xlsx_mw(ws,"A16:X16","Procedimiento - Control de Energías Peligrosas / LOTO",fill=MODO3_COLOR_HEX,font_color=MODO3_FONT_HEX,bold=True,size=8)
    _xlsx_mw(ws,"A17:X28","FOTO DEL PASO A PASO",fill="FFFFFF",font_color="94A3B8",bold=True,size=14)
    _xlsx_add_image(ws,ctx.get("photo_uri",""),"H17",max_width_px=520,max_height_px=210)

    # Tabla bloqueos — encabezado
    cur_row = 29
    _xlsx_mw(ws,f"A{cur_row}:X{cur_row}","Puntos de Bloqueo LOTO",fill=MODO3_COLOR_HEX,font_color=MODO3_FONT_HEX,bold=True,size=8)
    cur_row += 1
    # Cabecera columnas: A-B | C-F | G-J | K-N | O-R | S-X
    hdrs = [("A","B","N°"),("C","F","Fuente de Energía y Magnitud"),("G","J","Ubicación"),("K","N","Acción"),("O","R","Verificación"),("S","X","Dispositivo de Bloqueo")]
    for cs, ce, ht in hdrs:
        _xlsx_mw(ws,f"{cs}{cur_row}:{ce}{cur_row}",ht,fill=MODO3_COLOR_HEX,font_color=MODO3_FONT_HEX,bold=True,size=7)
    cur_row += 1

    for i, p in enumerate(lockpoints, 1):
        bg, fc = _energia_colors(p.get("energia",""))
        row_fill = "FFF7ED" if i % 2 == 0 else "FFFFFF"
        _xlsx_mw(ws,f"A{cur_row}:B{cur_row}",str(i),fill="FFF7ED",font_color="0F172A",bold=True,size=12)
        _xlsx_mw(ws,f"C{cur_row}:F{cur_row}",f"{p.get('energia','')}\n{p.get('magnitud','')}",fill=bg,font_color=fc,bold=True,size=7)
        _xlsx_mw(ws,f"G{cur_row}:J{cur_row}",p.get("ubicacion",""),fill=row_fill,size=7,align="left")
        _xlsx_mw(ws,f"K{cur_row}:N{cur_row}",p.get("accion",""),fill=row_fill,size=7,align="left")
        _xlsx_mw(ws,f"O{cur_row}:R{cur_row}",p.get("verificacion",""),fill=row_fill,size=7,align="left")
        _xlsx_mw(ws,f"S{cur_row}:X{cur_row}",p.get("dispositivo",""),fill=row_fill,size=7,align="left")
        ws.row_dimensions[cur_row].height = 28
        cur_row += 1

    # Procedimiento
    _xlsx_mw(ws,f"A{cur_row}:L{cur_row}","Acción",fill="3B3B3B",font_color="FFFFFF",bold=True,size=7); 
    _xlsx_mw(ws,f"M{cur_row}:X{cur_row}","Verificación",fill="3B3B3B",font_color="FFFFFF",bold=True,size=7); cur_row+=1
    ar, vr = cur_row, cur_row
    _xlsx_mw(ws,f"A{ar}:L{ar+11}","\n\n".join(_ACCION_MODO3),fill="FFFFFF",size=7.0,align="left")
    _xlsx_mw(ws,f"M{vr}:X{vr+11}","\n\n".join(_VERIFICACION_MODO3),fill="FFFFFF",size=7.0,align="left")
    for r in range(ar, ar+12): ws.row_dimensions[r].height = 22
    cur_row = ar + 12

    # Leyendas
    _xlsx_mw(ws,f"A{cur_row}:X{cur_row}","Clasificación de Energías Peligrosas",fill=MODO3_COLOR_HEX,font_color=MODO3_FONT_HEX,bold=True,size=8); cur_row+=1
    from openpyxl.utils import get_column_letter as gcl
    xl_en = [
        (1,4,"E: Eléctrica","FFFFFF","000000",None,None),(5,8,"N: Neumática","FFFFFF","0284C7",None,None),
        (9,12,"AM: Amoníaco","0F172A","D9D9D9","F59E0B",[(15,26),(42,53),(68,79)]),
        (13,16,"T: Térmica","FFFFFF","DC2626",None,None),(17,20,"H: Hidráulica","FFFFFF","7C3AED",None,None),
        (21,24,"P: Potencial","0F172A","FFF200","050505",[(22,34),(62,74)]),
    ]
    xl_en2 = [
        (1,4,"Q: Química","0F172A","FFF200",None,None),(5,8,"V: Vapor","111827","F59E0B",None,None),
        (9,12,"A: Agua","FFFFFF","16A34A",None,None),
        (13,16,"SC: Soda Cáustica","0F172A","D9D9D9","F59E0B",[(25,37),(60,72)]),
        (17,20,"Oz: Ozono","0F172A","BAE6FD",None,None),(21,24,"GC: Gas Carbónico","111827","C7D2FE",None,None),
    ]
    from openpyxl.styles import Alignment,Border,Font,PatternFill,Side
    side = Side(style="thin",color="111827"); brd = Border(left=side,right=side,top=side,bottom=side)
    for row_offset, items in enumerate([xl_en, xl_en2]):
        r = cur_row + row_offset
        ws.row_dimensions[r].height = 24
        for cs,ce,lbl,fc,bg,_sc,_pos in items:
            ws.merge_cells(f"{gcl(cs)}{r}:{gcl(ce)}{r}")
            cell = ws.cell(row=r,column=cs); cell.value = lbl
            cell.fill = PatternFill("solid",fgColor=bg); cell.font = Font(name="Bahnschrift",size=7,bold=True,color=fc)
            cell.alignment = Alignment(horizontal="center",vertical="center",wrap_text=True)
            for c in range(cs,ce+1): ws.cell(row=r,column=c).border = brd
    cur_row += 2

    _xlsx_mw(ws,f"A{cur_row}:X{cur_row}","Clasificación de Candados según sector y función",fill=MODO3_COLOR_HEX,font_color=MODO3_FONT_HEX,bold=True,size=8); cur_row+=1
    for rng2,lbl,fill,fc in [(f"A{cur_row}:E{cur_row+1}","Mantenimiento\nIndustrial","EF0000","FFFFFF"),
                              (f"F{cur_row}:J{cur_row+1}","Calidad","FFF200","0F172A"),
                              (f"K{cur_row}:N{cur_row+1}","Producción","16A34A","FFFFFF"),
                              (f"O{cur_row}:S{cur_row+1}","Mantenimiento Edilicio\nContratistas","0F7DBD","FFFFFF"),
                              (f"T{cur_row}:X{cur_row+1}","Supervisor MMTO Industrial\n(Bloqueo Departamental)","050505","FFFFFF")]:
        _xlsx_mw(ws,rng2,lbl,fill=fill,font_color=fc,bold=True,size=7)
        for r in range(cur_row,cur_row+2): ws.row_dimensions[r].height = 24
    cur_row += 2

    _xlsx_mw(ws,f"A{cur_row}:X{cur_row+6}","LSR",fill="FFFFFF",font_color="94A3B8",bold=True,size=14)
    _xlsx_add_image(ws,lsr_uri,f"D{cur_row}",max_width_px=600,max_height_px=130)
    cur_row += 7

    _xlsx_mw(ws,f"A{cur_row}:L{cur_row+2}",f"Elaborado por: {elaborado_por or '-'}\nPuesto: {puesto_elaborado or '-'} · Fecha: {fecha_firma_txt}",fill="F8FAFC",size=7)
    _xlsx_mw(ws,f"M{cur_row}:X{cur_row+2}",f"Aprobado por: {aprobado_por or '-'}\nPuesto: {puesto_aprobado or '-'} · Fecha: {fecha_firma_txt}",fill="F8FAFC",size=7)

    ws.print_area = f"A1:X{cur_row+2}"
    out = io.BytesIO(); wb.save(out); return out.getvalue()

# ============================================================
# INTERFAZ STREAMLIT
# ============================================================

st.markdown("""<style>
html,body,[class*="css"],.stApp,input,textarea,button{
    font-family:Bahnschrift,'Bahnschrift SemiCondensed','Arial Narrow',Arial,sans-serif !important;}
</style>""", unsafe_allow_html=True)

st.title("GENERADOR DE PROCEDIMIENTOS MODO 3 / LOTO")
st.markdown("""<div class="soft-note">
Importá el borrador JSON. Esta versión genera el procedimiento <strong>Modo 3 / LOTO</strong>
con tabla de puntos de bloqueo dinámica basada en las energías del JSON.
Podés agregar, editar o eliminar puntos antes de generar.
</div>""", unsafe_allow_html=True)

with st.sidebar:
    st.subheader("Configuración documental")
    codigo       = st.text_input("Código", value="LOTO-M3-001")
    revision     = st.text_input("Revisión", value="00")
    fecha        = st.date_input("Fecha", value=datetime.date.today(), format="DD/MM/YYYY")
    organizacion = st.text_input("Organización / logo textual", value="ARCA CONTINENTAL")
    logo_file    = st.file_uploader("Logo (reemplaza arca.png)", type=["png","jpg","jpeg","webp"])
    st.divider()
    st.subheader("Firmas")
    elaborado_por   = st.text_input("Elaborado por", value="")
    puesto_elaborado= st.text_input("Puesto de quien elabora", value="")
    aprobado_por    = st.text_input("Aprobado por", value="")
    puesto_aprobado = st.text_input("Puesto de quien aprueba", value="")
    fecha_firma     = st.date_input("Fecha de firmas", value=fecha, format="DD/MM/YYYY")

uploaded_json = st.file_uploader("Importar borrador JSON", type=["json"])
if uploaded_json is None:
    st.info("Cargá un borrador JSON para generar el procedimiento.")
    st.stop()

try:
    payload = json.loads(uploaded_json.getvalue().decode("utf-8"))
except Exception as exc:
    st.error(f"No se pudo leer el JSON: {exc}"); st.stop()

ctx_detected = extract_procedure_context(payload)

st.subheader("Datos importados editables")
c1, c2, c3 = st.columns(3)
with c1:
    negocio_edit   = st.text_input("Negocio",    value=_normalize(ctx_detected.get("negocio")),    key="edit_negocio")
    sitio_edit     = st.text_input("Sitio",      value=_normalize(ctx_detected.get("sitio")),      key="edit_sitio")
    area_edit      = st.text_input("Área",       value=_normalize(ctx_detected.get("area")),       key="edit_area")
with c2:
    linea_edit     = st.text_input("Línea",      value=_normalize(ctx_detected.get("linea")),      key="edit_linea")
    equipo_edit    = st.text_input("Equipo",     value=_normalize(ctx_detected.get("equipo")),     key="edit_equipo")
    fabricante_edit= st.text_input("Fabricante", value=_normalize(ctx_detected.get("fabricante")), key="edit_fabricante")
with c3:
    modelo_edit    = st.text_input("Modelo",     value=_normalize(ctx_detected.get("modelo")),     key="edit_modelo")
    anio_edit      = st.text_input("Año",        value=_normalize(ctx_detected.get("anio")),       key="edit_anio")
    modo_final_edit= st.text_input("Modo final", value=_normalize(ctx_detected.get("modo_final")), key="edit_modo_final")

st.subheader("Evaluación de riesgos de referencia")
ec1, ec2 = st.columns(2)
with ec1:
    eval_riesgos_codigo = st.text_input("Código de la evaluación",
        value=_normalize(ctx_detected.get("codigo_documento")), key="eval_cod",
        help="Campo 'codigo_documento' del JSON.")
with ec2:
    eval_riesgos_fecha = st.text_input("Fecha de la evaluación",
        value=_normalize(ctx_detected.get("fecha_documento")), key="eval_fec",
        help="Campo 'fecha_documento' del JSON.")

st.subheader("Personal")
pc1, pc2 = st.columns(2)
with pc1:
    personal_afectado  = st.text_area("Personal afectado",   value="Técnicos del Sector de Mantenimiento.\nContratistas.", height=80, key="pa")
with pc2:
    personal_autorizado= st.text_area("Personal autorizado", value="Técnicos del Sector Mantenimiento.\nContratistas.",   height=80, key="pu")

st.subheader("Tareas aplicables")
tasks_detected = ctx_detected.get("tareas") or []
tasks_text = st.text_area("Tareas (predefinidas + manuales combinadas)", value="\n".join(tasks_detected),
    height=120, help="Una tarea por línea. Se combinan 'tareas_predefinidas' y 'tareas_manuales' del JSON.")

# ── PUNTOS DE BLOQUEO ──────────────────────────────────────────────────────
st.subheader("Puntos de bloqueo")
st.caption("Se generan automáticamente desde las energías del JSON. Podés editar cada punto, agregar duplicados de la misma energía o añadir nuevos manualmente.")

energias_json = ctx_detected.get("energias") or {}

# Inicializar en session_state
if "lockpoints" not in st.session_state or st.button("↺ Resetear desde JSON", help="Restaura los puntos originales del JSON"):
    st.session_state.lockpoints = energias_to_lockpoints(energias_json)

# Lista de nombres de energía para el selectbox
ENERGIA_NOMBRES = ["Eléctrica","Neumática","Térmica","Hidráulica","Potencial","Química","Vapor","Agua","Amoníaco","Soda Cáustica","Ozono","Gas Carbónico","Otra"]

# Botón agregar
if st.button("＋ Agregar punto de bloqueo"):
    st.session_state.lockpoints.append({"energia":"Eléctrica","magnitud":"","ubicacion":"","accion":"","verificacion":"","dispositivo":""})

lp_to_delete = []
for i, pt in enumerate(st.session_state.lockpoints):
    bg, fc = _energia_colors(pt.get("energia",""))
    header_color = f"#{bg}"
    st.markdown(f"""<div style='background:{header_color};color:#{fc};font-weight:900;font-size:12px;
        padding:5px 10px;border-radius:6px 6px 0 0;margin-top:8px;'>
        Punto {i+1} — {pt.get('energia','')} {pt.get('magnitud','')}
    </div>""", unsafe_allow_html=True)

    col_e, col_m, col_del = st.columns([2, 2, 0.5])
    with col_e:
        energia_options = ENERGIA_NOMBRES
        cur_val = pt.get("energia","Eléctrica")
        if cur_val not in energia_options:
            energia_options = [cur_val] + energia_options
        idx = energia_options.index(cur_val) if cur_val in energia_options else 0
        pt["energia"] = st.selectbox(f"Energía #{i+1}", energia_options, index=idx, key=f"en_{i}", label_visibility="collapsed")
    with col_m:
        pt["magnitud"] = st.text_input(f"Magnitud #{i+1}", value=pt.get("magnitud",""), placeholder="Magnitud / descripción", key=f"mag_{i}", label_visibility="collapsed")
    with col_del:
        if st.button("🗑", key=f"del_{i}", help="Eliminar punto"):
            lp_to_delete.append(i)

    col_ub, col_ac = st.columns(2)
    with col_ub:
        pt["ubicacion"] = st.text_input(f"Ubicación #{i+1}", value=pt.get("ubicacion",""), placeholder="Ej: Tablero eléctrico principal", key=f"ub_{i}")
    with col_ac:
        pt["accion"] = st.text_input(f"Acción #{i+1}", value=pt.get("accion",""), placeholder="Ej: Abrir y bloquear disyuntor Q1", key=f"ac_{i}")

    col_ver, col_dis = st.columns(2)
    with col_ver:
        pt["verificacion"] = st.text_input(f"Verificación #{i+1}", value=pt.get("verificacion",""), placeholder="Ej: Medir tensión con tester", key=f"vr_{i}")
    with col_dis:
        pt["dispositivo"] = st.text_input(f"Dispositivo #{i+1}", value=pt.get("dispositivo",""), placeholder="Ej: Candado rojo + grapa", key=f"dv_{i}")

for i in reversed(lp_to_delete):
    st.session_state.lockpoints.pop(i)

lockpoints = st.session_state.lockpoints

st.markdown(f"**Total de puntos de bloqueo: {len(lockpoints)}**")

# Foto
st.subheader("Foto del paso a paso")
step_photo = st.file_uploader("Subir foto", type=["png","jpg","jpeg","webp"])
photo_uri  = _uploaded_file_data_uri(step_photo)
if step_photo:
    st.image(step_photo, caption="Foto cargada", use_container_width=True)
else:
    st.caption("Sin foto: se mostrará recuadro reservado.")

# Contexto final
ctx = dict(ctx_detected)
ctx.update({
    "negocio": negocio_edit, "sitio": sitio_edit, "area": area_edit, "linea": linea_edit,
    "equipo": equipo_edit, "fabricante": fabricante_edit, "modelo": modelo_edit, "anio": anio_edit,
    "modo_final": modo_final_edit, "tareas": _split_tasks(tasks_text), "photo_uri": photo_uri,
})
modo_final = ctx.get("modo_final")

# Métricas
m1, m2, m3, m4 = st.columns(4)
m1.metric("Modo final",      _normalize(modo_final) or "Sin dato")
m2.metric("Equipo",          (_normalize(ctx.get("equipo")) or "Sin dato")[:28])
m3.metric("Tareas",          len(ctx.get("tareas") or []))
m4.metric("Puntos bloqueo",  len(lockpoints))

if not _mode_is_modo_3(modo_final):
    st.markdown(f"""<div class="status-warn">
        El JSON no corresponde a <strong>Modo 3</strong>. Modo detectado: <strong>{_html_esc(modo_final or 'sin dato')}</strong>.
    </div>""", unsafe_allow_html=True)
    with st.expander("Ver datos detectados"): st.json({"ctx": ctx, "computed": payload.get("computed",{})})
    st.stop()

logo_uri = _logo_data_uri(logo_file)
lsr_uri  = _lsr_data_uri()

procedure_html = build_modo_3_html(
    ctx, lockpoints, codigo=codigo, revision=revision, fecha=fecha, organizacion=organizacion,
    logo_uri=logo_uri, lsr_uri=lsr_uri, personal_afectado=personal_afectado,
    personal_autorizado=personal_autorizado, elaborado_por=elaborado_por, aprobado_por=aprobado_por,
    puesto_elaborado=puesto_elaborado, puesto_aprobado=puesto_aprobado, fecha_firma=fecha_firma,
    eval_riesgos_codigo=eval_riesgos_codigo, eval_riesgos_fecha=eval_riesgos_fecha,
)

st.markdown("""<div class="status-ok">Borrador compatible con <strong>Modo 3 / LOTO</strong>. Vista previa y descargas habilitadas.</div>""", unsafe_allow_html=True)

st.subheader("Vista previa del procedimiento")
components.html(procedure_html, height=1100, scrolling=True)

stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
base  = f"Procedimiento_Modo_3_{stamp}"
shared_kwargs = dict(ctx=ctx, lockpoints=lockpoints, codigo=codigo, revision=revision, fecha=fecha,
    organizacion=organizacion, logo_uri=logo_uri, lsr_uri=lsr_uri,
    personal_afectado=personal_afectado, personal_autorizado=personal_autorizado,
    elaborado_por=elaborado_por, aprobado_por=aprobado_por,
    puesto_elaborado=puesto_elaborado, puesto_aprobado=puesto_aprobado,
    fecha_firma=fecha_firma, eval_riesgos_codigo=eval_riesgos_codigo, eval_riesgos_fecha=eval_riesgos_fecha)

col_w, col_x, col_h, col_j = st.columns(4)
with col_w:
    st.download_button("Descargar Word DOCX", data=html_to_word_bytes(**shared_kwargs),
        file_name=f"{base}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True)
with col_x:
    try:
        st.download_button("Descargar Excel XLSX", data=build_modo_3_excel_bytes(**shared_kwargs),
            file_name=f"{base}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True)
    except Exception as exc:
        st.warning(f"Excel no disponible: {exc}")
with col_h:
    st.download_button("Descargar HTML", data=procedure_html.encode("utf-8"),
        file_name=f"{base}.html", mime="text/html", use_container_width=True)
with col_j:
    st.download_button("Descargar JSON", data=json.dumps(payload, ensure_ascii=False, indent=4).encode("utf-8"),
        file_name=f"Borrador_{stamp}.json", mime="application/json", use_container_width=True)

with st.expander("Datos para generación"):
    st.json({"equipo": ctx.get("equipo"), "modo_final": ctx.get("modo_final"),
             "tareas": ctx.get("tareas"), "n_bloqueos": len(lockpoints),
             "puntos": lockpoints, "eval_cod": eval_riesgos_codigo, "eval_fec": eval_riesgos_fecha})
